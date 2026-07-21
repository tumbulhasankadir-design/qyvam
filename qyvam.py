import streamlit as st
import sqlite3
import io
import os
import pytz
import random
from datetime import datetime
from docx import Document
from docx.shared import Inches
import pandas as pd
import plotly.express as px

# ==============================================================================
# SİSTEM AYARLARI VE GİZLİ KASA
# ==============================================================================
st.set_page_config(page_title="Qyvam | Siber Uzay", layout="wide", initial_sidebar_state="expanded")

try:
    from openai import OpenAI
    OPENROUTER_KEY = st.secrets["OPENROUTER_KEY"]
except:
    OPENROUTER_KEY = ""

try:
    from mufredat import MUFREDAT
except ImportError:
    MUFREDAT = {}
    st.error("[ SİSTEM UYARISI ]: mufredat.py dosyası bulunamadı.")

DB_YOLU = 'qyvam_siber.db'

# Oturum Değişkenleri
if 'aktif_sayfa' not in st.session_state: st.session_state.aktif_sayfa = "Ana Sayfa"
if 'veli_kadi' not in st.session_state: st.session_state.veli_kadi = None
if 'aktif_cocuk_id' not in st.session_state: st.session_state.aktif_cocuk_id = None
if 'aktif_cocuk_isim' not in st.session_state: st.session_state.aktif_cocuk_isim = ""

# ==============================================================================
# VERİTABANI MOTORU
# ==============================================================================
def init_db():
    conn = sqlite3.connect(DB_YOLU)
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS Veliler (kullanici_adi TEXT PRIMARY KEY, sifre TEXT)''')
    c.execute('''CREATE TABLE IF NOT EXISTS Cocuklar (id INTEGER PRIMARY KEY AUTOINCREMENT, isim TEXT, mevcut_adim INTEGER, veli_kadi TEXT)''')
    try:
        c.execute("ALTER TABLE Cocuklar ADD COLUMN veli_kadi TEXT DEFAULT 'Kurucu'")
    except:
        pass
    c.execute('''CREATE TABLE IF NOT EXISTS Gorevler (id INTEGER PRIMARY KEY AUTOINCREMENT, cocuk_id INTEGER, gorev_adi TEXT, tefekkur TEXT, durum TEXT, puan INTEGER, veli_notu TEXT)''')
    c.execute('''CREATE TABLE IF NOT EXISTS ozel_beratlar (id INTEGER PRIMARY KEY AUTOINCREMENT, cocuk_id INTEGER, berat_adi TEXT, berat_aciklama TEXT)''')
    c.execute('''CREATE TABLE IF NOT EXISTS Takvim (id INTEGER PRIMARY KEY AUTOINCREMENT, veli_kadi TEXT, cocuk_id INTEGER, tarih TEXT, baslik TEXT, detay TEXT)''')
    conn.commit()
    conn.close()

init_db()

# --- HAYALET VERİ TEMİZLİĞİ (Sadece 1 kez çalışması yeterli) ---
conn_temizle = sqlite3.connect(DB_YOLU)
conn_temizle.cursor().execute("DELETE FROM Cocuklar WHERE isim = 'Yusuf' OR veli_kadi = 'Kurucu'")
conn_temizle.commit()
conn_temizle.close()

# --- Temel Veritabanı Fonksiyonları ---
def veli_kaydol(kadi, sifre):
    conn = sqlite3.connect(DB_YOLU)
    c = conn.cursor()
    try:
        c.execute("INSERT INTO Veliler (kullanici_adi, sifre) VALUES (?, ?)", (kadi, sifre))
        conn.commit()
        basari = True
    except:
        basari = False
    conn.close()
    return basari

def veli_giris_yap(kadi, sifre):
    conn = sqlite3.connect(DB_YOLU)
    c = conn.cursor()
    c.execute("SELECT * FROM Veliler WHERE kullanici_adi=? AND sifre=?", (kadi, sifre))
    kullanici = c.fetchone()
    conn.close()
    return kullanici is not None

def cocuk_ekle(isim, veli_kadi):
    conn = sqlite3.connect(DB_YOLU)
    c = conn.cursor()
    c.execute("INSERT INTO Cocuklar (isim, mevcut_adim, veli_kadi) VALUES (?, ?, ?)", (isim, 1, veli_kadi))
    conn.commit()
    conn.close()

def cocuk_sil(cocuk_id):
    conn = sqlite3.connect(DB_YOLU)
    c = conn.cursor()
    c.execute("DELETE FROM Cocuklar WHERE id = ?", (cocuk_id,))
    c.execute("DELETE FROM Gorevler WHERE cocuk_id = ?", (cocuk_id,))
    c.execute("DELETE FROM ozel_beratlar WHERE cocuk_id = ?", (cocuk_id,))
    c.execute("DELETE FROM Takvim WHERE cocuk_id = ?", (cocuk_id,))
    conn.commit()
    conn.close()

def cocuk_adim_guncelle(cocuk_id, yeni_adim):
    conn = sqlite3.connect(DB_YOLU)
    c = conn.cursor()
    c.execute("UPDATE Cocuklar SET mevcut_adim = ? WHERE id = ?", (yeni_adim, cocuk_id))
    conn.commit()
    conn.close()

def cocuklari_getir(veli_kadi):
    conn = sqlite3.connect(DB_YOLU)
    c = conn.cursor()
    c.execute("SELECT id, isim, mevcut_adim FROM Cocuklar WHERE veli_kadi=?", (veli_kadi,))
    data = c.fetchall()
    conn.close()
    return data

def tum_cocuklari_getir():
    conn = sqlite3.connect(DB_YOLU)
    c = conn.cursor()
    c.execute("SELECT id, isim, veli_kadi FROM Cocuklar")
    data = c.fetchall()
    conn.close()
    return data

def cocuk_bilgisi_getir(cocuk_id):
    conn = sqlite3.connect(DB_YOLU)
    imlec = conn.cursor()
    imlec.execute('SELECT isim, mevcut_adim FROM Cocuklar WHERE id = ?', (cocuk_id,))
    sonuc = imlec.fetchone()
    conn.close()
    return sonuc

# --- Görev ve Onay Fonksiyonları ---
def onay_bekleyenleri_getir(veli_kadi):
    conn = sqlite3.connect(DB_YOLU)
    imlec = conn.cursor()
    imlec.execute('''SELECT G.id, C.isim, G.gorev_adi, G.tefekkur, C.id, C.mevcut_adim FROM Gorevler G JOIN Cocuklar C ON G.cocuk_id = C.id WHERE G.durum = 'Veli Onayı Bekliyor' AND C.veli_kadi = ?''', (veli_kadi,))
    liste = imlec.fetchall()
    conn.close()
    return liste

def ozel_beratlari_getir(cocuk_id):
    conn = sqlite3.connect(DB_YOLU)
    c = conn.cursor()
    c.execute("SELECT berat_adi, berat_aciklama FROM ozel_beratlar WHERE cocuk_id=?", (cocuk_id,))
    veriler = c.fetchall()
    conn.close()
    return veriler

def ozel_berat_ekle(cocuk_id, berat_adi, berat_aciklama):
    conn = sqlite3.connect(DB_YOLU)
    c = conn.cursor()
    c.execute("INSERT INTO ozel_beratlar (cocuk_id, berat_adi, berat_aciklama) VALUES (?, ?, ?)", (cocuk_id, berat_adi, berat_aciklama))
    conn.commit()
    conn.close()

def bekleyen_gorev_kontrol(cocuk_id):
    conn = sqlite3.connect(DB_YOLU)
    imlec = conn.cursor()
    imlec.execute("SELECT id FROM Gorevler WHERE cocuk_id=? AND durum='Veli Onayı Bekliyor'", (cocuk_id,))
    sonuc = imlec.fetchone()
    conn.close()
    return sonuc is not None

def cocuk_veri_gonder(cocuk_id, gorev_adi, tefekkur_cevabi):
    conn = sqlite3.connect(DB_YOLU)
    imlec = conn.cursor()
    imlec.execute("INSERT INTO Gorevler (cocuk_id, gorev_adi, tefekkur, durum) VALUES (?, ?, ?, 'Veli Onayı Bekliyor')", (cocuk_id, gorev_adi, tefekkur_cevabi))
    conn.commit()
    conn.close()

def veri_onayla(islem_id, cocuk_id):
    conn = sqlite3.connect(DB_YOLU)
    imlec = conn.cursor()
    imlec.execute("UPDATE Gorevler SET durum = 'Onaylandı' WHERE id = ?", (islem_id,))
    imlec.execute("UPDATE Cocuklar SET mevcut_adim = mevcut_adim + 1 WHERE id = ?", (cocuk_id,))
    conn.commit()
    conn.close()

# --- Takvim Fonksiyonları ---
def ajanda_ekle(veli_kadi, cocuk_id, tarih, baslik, detay):
    conn = sqlite3.connect(DB_YOLU)
    c = conn.cursor()
    c.execute("INSERT INTO Takvim (veli_kadi, cocuk_id, tarih, baslik, detay) VALUES (?, ?, ?, ?, ?)", (veli_kadi, cocuk_id, str(tarih), baslik, detay))
    conn.commit()
    conn.close()

def ajanda_getir(veli_kadi):
    conn = sqlite3.connect(DB_YOLU)
    c = conn.cursor()
    c.execute('''SELECT T.id, C.isim, T.tarih, T.baslik, T.detay FROM Takvim T JOIN Cocuklar C ON T.cocuk_id = C.id WHERE T.veli_kadi=? ORDER BY T.tarih ASC''', (veli_kadi,))
    veriler = c.fetchall()
    conn.close()
    return veriler

def cocuk_ajanda_getir(cocuk_id):
    conn = sqlite3.connect(DB_YOLU)
    c = conn.cursor()
    bugun = datetime.now().strftime("%Y-%m-%d")
    c.execute("SELECT tarih, baslik, detay FROM Takvim WHERE cocuk_id=? AND tarih >= ? ORDER BY tarih ASC", (cocuk_id, bugun))
    veriler = c.fetchall()
    conn.close()
    return veriler

def ajanda_sil(kayit_id):
    conn = sqlite3.connect(DB_YOLU)
    c = conn.cursor()
    c.execute("DELETE FROM Takvim WHERE id=?", (kayit_id,))
    conn.commit()
    conn.close()

# ==============================================================================
# YARDIMCI & RAPORLAMA FONKSİYONLARI
# ==============================================================================
def kazanilan_sertifikalari_hesapla(adim):
    sertifikalar = []
    if adim > 14: sertifikalar.append(("[ KÖKLER YETKİNLİK BERATI ]", "Bedenine ve mekanına gösterdiği irade ile kendi şahsi gayreti ve rehberinin onayıyla kazanılmıştır."))
    if adim > 28: sertifikalar.append(("[ BAĞLAR YETKİNLİK BERATI ]", "Ailesine ve çevresine gösterdiği ikram bilinciyle kazanılmıştır."))
    if adim > 42: sertifikalar.append(("[ PUSULA YETKİNLİK BERATI ]", "Zaman ve dijital irade yönetimindeki kararlılığıyla kazanılmıştır."))
    if adim > 56: sertifikalar.append(("[ AYNALAR YETKİNLİK BERATI ]", "Öz-farkındalık ve derin empati yeteneğiyle tescillenmiştir."))
    if adim > 70: sertifikalar.append(("[ ÇARKLAR YETKİNLİK BERATI ]", "Üretim ve finansal denge konularındaki olgunluğuyla kazanılmıştır."))
    if adim > 84: sertifikalar.append(("[ KÖPRÜLER YETKİNLİK BERATI ]", "Topluma ve doğaya olan merhametli yaklaşımıyla kazanılmıştır."))
    if adim > 99: sertifikalar.append(("[ KANATLAR YETKİNLİK BERATI ]", "İhsan makamında, örnek bir şahsiyet olmaya dair üstün yetkinlik belgesidir."))
    return sertifikalar
    
def word_raporu_olustur(cocuk_ismi, gorev_adi, veli_notu, fotograf_bytes=None):
    doc = Document()
    doc.add_heading(f'Qyvam - Görev Raporu: {cocuk_ismi}', 0)
    doc.add_paragraph("Görev İçeriği:").bold = True
    doc.add_paragraph(gorev_adi)
    doc.add_paragraph("Rehber (Veli) Gözlemi:").bold = True
    doc.add_paragraph(veli_notu)
    if fotograf_bytes:
        doc.add_paragraph("Görevin Görsel Kanıtı:").bold = True
        foto_io = io.BytesIO(fotograf_bytes)
        doc.add_picture(foto_io, width=Inches(5.0))
    byte_io = io.BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)
    return byte_io

def ai_cevap_uret(soru, mevcut_adim, rol="veli", cocuk_isim=""):
    if not OPENROUTER_KEY:
        return "[ SİSTEM UYARISI ]: API Anahtarı bulunamadı. Yapay zeka çevrimdışı."
    faz = MUFREDAT.get(mevcut_adim, {}).get("faz", "Bilinmeyen Faz")
    if rol == "veli":
        prompt = f"Sen Qyvam adında uzman bir pedagojik asistansın. Veliye cevap veriyorsun. Çocuğun mevcut aşaması: {faz}. Soru: {soru}"
    else:
        prompt = f"Senin adın Qyman. Zeki, siber fütüristik bir dijital ikiz rehberisin. Karşındaki arkadaşının adı {cocuk_isim}. Cümlelerin kısa, cesaretlendirici olsun. Emojileri asla kullanma. Gelişim aşaması: {faz}. Soru: {soru}"
    try:
        client = OpenAI(base_url="https://openrouter.ai/api/v1", api_key=OPENROUTER_KEY)
        response = client.chat.completions.create(model="deepseek/deepseek-v4-flash", messages=[{"role": "user", "content": prompt}])
        return response.choices[0].message.content
    except Exception as e:
        return f"[ SİSTEM HATASI ]: Bağlantı kurulamadı. Detay: {str(e)}"

# ==============================================================================
# ARAYÜZ TASARIMI (WEB SİTESİ ESTETİĞİ & TAKVİM CSS)
# ==============================================================================
st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;600;700&family=Plus+Jakarta+Sans:wght@400;500;600;700&display=swap');
    
    html, body, [class*="css"] { font-family: 'Plus Jakarta Sans', sans-serif; background-color: #f8fafc !important; color: #1e293b !important; }
    h1, h2, h3, h4, h5 { font-family: 'Outfit', sans-serif; font-weight: 700; color: #0f172a !important; letter-spacing: -0.5px; }
    
    .glass-box { background: #ffffff; border: 1px solid #e2e8f0; border-radius: 20px; padding: 30px; margin-bottom: 25px; box-shadow: 0 10px 25px -5px rgba(0, 0, 0, 0.05); transition: transform 0.2s ease; }
    .glass-box:hover { transform: translateY(-2px); box-shadow: 0 20px 25px -5px rgba(0, 0, 0, 0.05); }
    
    .glass-task { background: linear-gradient(145deg, #ffffff, #fdf4ff); border-left: 6px solid #d946ef; border-radius: 16px; padding: 25px; margin-bottom: 20px; box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.05); }
    
    .agenda-card { background: #ffffff; border-left: 5px solid #6366f1; border-radius: 12px; padding: 15px 20px; margin-bottom: 15px; box-shadow: 0 2px 10px rgba(99, 102, 241, 0.08); display: flex; flex-direction: column; }
    .agenda-date { color: #6366f1; font-weight: 700; font-size: 0.9rem; letter-spacing: 1px; text-transform: uppercase; margin-bottom: 5px; }
    .agenda-title { font-size: 1.15rem; font-weight: 700; color: #1e293b; margin-bottom: 5px; }
    .agenda-detail { font-size: 0.95rem; color: #64748b; }
    
    .neon-alert { background: linear-gradient(90deg, #ec4899, #8b5cf6); border-radius: 12px; padding: 15px 20px; color: white; margin-bottom: 25px; box-shadow: 0 4px 15px rgba(236, 72, 153, 0.3); }
    
    .top-bar { background: rgba(255, 255, 255, 0.8); backdrop-filter: blur(12px); border-bottom: 1px solid #e2e8f0; padding: 15px; margin-top: -60px; margin-bottom: 40px; border-radius: 0 0 24px 24px; }
    .qyman-hud { background: linear-gradient(135deg, #eff6ff 0%, #e0e7ff 100%); border: 1px solid #bfdbfe; border-radius: 16px; padding: 25px; color: #1e3a8a; line-height: 1.7; margin-bottom: 25px; }
    .neon-text { background: linear-gradient(to right, #4f46e5, #9333ea); -webkit-background-clip: text; -webkit-text-fill-color: transparent; font-weight: 800; }
    
    .stButton>button { border-radius: 14px; font-family: 'Plus Jakarta Sans', sans-serif; font-size: 1rem !important; font-weight: 600; border: none; background: linear-gradient(135deg, #6366f1 0%, #4f46e5 100%); color: white !important; transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1); padding: 12px 24px; box-shadow: 0 4px 6px -1px rgba(99, 102, 241, 0.4); }
    .stButton>button:hover { background: linear-gradient(135deg, #4f46e5 0%, #4338ca 100%); transform: translateY(-2px); }
    
    .hero-title { text-align: center; font-size: 5rem; background: linear-gradient(to right, #3b82f6, #8b5cf6, #ec4899); -webkit-background-clip: text; -webkit-text-fill-color: transparent; margin-bottom: 5px; font-family: 'Outfit', sans-serif; font-weight: 900; letter-spacing: -2px; }
    .hero-subtitle { text-align: center; color: #64748b; font-size: 1.3rem; font-weight: 500; margin-bottom: 40px; }
    
    .stTabs [data-baseweb="tab-list"] { gap: 8px; background-color: #f1f5f9; border-radius: 12px; padding: 6px; }
    .stTabs [data-baseweb="tab"] { border-radius: 8px; padding: 10px 15px; color: #64748b; font-weight: 600; }
    .stTabs [aria-selected="true"] { background-color: #ffffff !important; color: #4f46e5 !important; box-shadow: 0 2px 4px rgba(0,0,0,0.05); }
    </style>
""", unsafe_allow_html=True)

# ==============================================================================
# ÜST DURUM ÇUBUĞU (NAVBAR - ŞIK VE TAM EKRAN)
# ==============================================================================
def ust_komuta_merkezi():
    st.markdown("""
    <style>
    /* Streamlit'in üst boşluğunu sıfırlıyoruz */
    .block-container { padding-top: 3.5rem !important; }
    
    /* Tam Ekran Üst Bar Arka Planı (Cam Efekti) */
    .modern-header-bg {
        position: absolute;
        top: 0;
        left: 50%;
        transform: translateX(-50%);
        width: 100vw;
        height: 90px;
        background: rgba(255, 255, 255, 0.85);
        backdrop-filter: blur(16px);
        -webkit-backdrop-filter: blur(16px);
        border-bottom: 1px solid rgba(226, 232, 240, 0.8);
        box-shadow: 0 4px 20px -2px rgba(0, 0, 0, 0.05);
        z-index: -1;
    }
    
    /* Navigasyon Yazı Stilleri */
    .nav-logo-text {
        font-family: 'Outfit', sans-serif;
        font-size: 1.6rem;
        font-weight: 900;
        background: linear-gradient(to right, #4f46e5, #9333ea);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        margin-bottom: 0;
        letter-spacing: 1px;
    }
    .nav-path { color: #64748b; font-weight: 500; font-size: 0.95rem; margin-top: -5px; }
    .nav-time { text-align: right; color: #475569; font-weight: 600; font-size: 1.05rem; padding-top: 15px; font-family: 'Plus Jakarta Sans', sans-serif; }
    
    /* Üst Menü Butonlarını Zarifleştirme */
    div[data-testid="column"]:nth-child(2) button {
        padding: 6px 16px !important;
        font-size: 0.95rem !important;
        border-radius: 20px !important;
        box-shadow: 0 4px 10px rgba(99, 102, 241, 0.2) !important;
        margin-top: 8px;
    }
    </style>
    <div class="modern-header-bg"></div>
    """, unsafe_allow_html=True)
    
    st.markdown('<div style="margin-top: -25px; margin-bottom: 40px;">', unsafe_allow_html=True)
    col_sol, col_orta, col_sag = st.columns([3, 2, 4])
    
    with col_sol:
        if st.session_state.aktif_sayfa == "Ana Sayfa": durum_metni = "Ana Sayfa"
        elif st.session_state.aktif_sayfa == "Veli_Panel": durum_metni = "Rehberlik Köşesi"
        elif st.session_state.aktif_sayfa == "Veli_Giris": durum_metni = "Rehber Girişi"
        else: durum_metni = "Gelişim Yolculuğu"
        
        st.markdown(f'<div class="nav-logo-text">✧ QYVAM</div>', unsafe_allow_html=True)
        st.markdown(f'<div class="nav-path">Eğitim Ekosistemi <span style="margin: 0 8px; color:#cbd5e1;">/</span> <b style="color: #4f46e5;">{durum_metni}</b></div>', unsafe_allow_html=True)
        
    with col_orta:
        if st.session_state.aktif_sayfa != "Ana Sayfa":
            if st.button("🏠 Ana Sayfa"):
                st.session_state.aktif_sayfa = "Ana Sayfa"
                st.session_state.aktif_cocuk_id = None
                st.rerun()
        if st.session_state.get("veli_kadi"):
            if st.session_state.aktif_sayfa != "Veli_Panel":
                if st.button("🛡️ Rehber Paneli"):
                    st.session_state.aktif_sayfa = "Veli_Panel"
                    st.rerun()
            else:
                if st.button("🚪 Çıkış Yap"):
                    st.session_state.veli_kadi = None
                    st.session_state.aktif_sayfa = "Ana Sayfa"
                    st.rerun()
        else:
            if st.session_state.aktif_sayfa not in ["Veli_Giris", "Veli_Panel"]:
                if st.button("🔐 Rehber Girişi"):
                    st.session_state.aktif_sayfa = "Veli_Giris"
                    st.rerun()
                    
    with col_sag:
        tz = pytz.timezone('Europe/Istanbul')
        simdi = datetime.now(tz)
        st.markdown(f'<div class="nav-time">{simdi.strftime("%d.%m.%Y")} <span style="color:#cbd5e1; margin:0 5px;">|</span> <span style="color:#6366f1;">{simdi.strftime("%H:%M")}</span></div>', unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

# ==============================================================================
# GELİŞİM RADARI (SOL MENÜ)
# ==============================================================================
def sol_radar_olustur():
    st.sidebar.markdown('<h2 class="neon-text" style="text-align:center; margin-bottom:0;">GELİŞİM RADARI</h2>', unsafe_allow_html=True)
    st.sidebar.markdown('<div style="text-align:center; margin-bottom: 30px; color:#64748b; font-size:0.9rem; font-weight:600;">Canlı Müfredat Takibi</div>', unsafe_allow_html=True)
    if not MUFREDAT: return
    faz_gruplari = {}
    for adim_no, detay in MUFREDAT.items():
        faz_adi = detay.get("faz", "Bilinmeyen Faz")
        ust_seviye = detay.get("ust_seviye", "Genel Eylem")
        alt_seviye = detay.get("alt_seviye", "Pekiştirme")
        if faz_adi not in faz_gruplari: faz_gruplari[faz_adi] = {}
        if ust_seviye not in faz_gruplari[faz_adi]: faz_gruplari[faz_adi][ust_seviye] = []
        faz_gruplari[faz_adi][ust_seviye].append((adim_no, alt_seviye))
    for faz_adi, ust_seviyeler in faz_gruplari.items():
        with st.sidebar.expander(f"📚 {faz_adi.upper()}"):
            for ust_adi, alt_liste in ust_seviyeler.items():
                st.markdown(f"<div style='color: #4f46e5; font-weight: 700; margin-top: 15px;'>{ust_adi}</div>", unsafe_allow_html=True)
                for adim_no, alt_adi in alt_liste:
                    st.markdown(f"<div style='color: #64748b; font-size: 0.9rem; padding-left: 10px; border-left: 2px solid #e2e8f0; margin-top: 5px;'>Adım {adim_no}: {alt_adi}</div>", unsafe_allow_html=True)

# ==============================================================================
# SAYFA 1 & 2: ANA KARŞILAMA VE GİRİŞ
# ==============================================================================
def ana_karsilama_ekrani():
    st.markdown('<div class="hero-title">Q Y V A M</div>', unsafe_allow_html=True)
    st.markdown('<div class="hero-subtitle">Bilişsel İklim ve Şahsiyet İnşası Serüveni</div>', unsafe_allow_html=True)
    col1, col2, col3 = st.columns([1, 1.5, 1])
    with col2:
        mevcut_videolar = [dosya for dosya in os.listdir() if dosya.startswith("qyman") and dosya.endswith(".mp4")]
        if mevcut_videolar: st.video(random.choice(mevcut_videolar), autoplay=True, loop=True, muted=True)
        elif os.path.exists("qyman.gif"): st.image("qyman.gif", use_container_width=True)
        elif os.path.exists("qyman.png"): st.image("qyman.png", use_container_width=True)
            
        st.markdown("""<div class="qyman-hud" style="text-align: center; margin-top: -20px; position: relative; z-index: 10;"><span style="font-size:1.3rem; font-weight:800; color:#4f46e5;">✨ Sisteme Hoş Geldin! ✨</span><br><br><b>Ben Qyman, senin dijital rehberinim.</b><br>Maceraya başlamak için aşağıdan profilini seç ve bağlantıyı kur.</div>""", unsafe_allow_html=True)
        
        cocuklar = tum_cocuklari_getir()
        if not cocuklar: st.warning("⚠️ Sistemde henüz kayıtlı bir profil yok. Lütfen sağ üstten 'Rehber Girişi' yaparak yeni bir hesap açın ve çocuğunuzu ekleyin.")
        else:
            st.markdown('<div class="glass-box"><h3 style="text-align:center; color:#0f172a; margin-bottom:20px;">[ BAĞLANTI MODÜLÜ ]</h3>', unsafe_allow_html=True)
            secenekler = [f"{c[1]} (Rehber: {c[2]})" for c in cocuklar]
            secilen_metin = st.selectbox("Kayıtlı Profiliniz:", secenekler, label_visibility="collapsed")
            st.markdown('<br>', unsafe_allow_html=True)
            if st.button("🚀 SİBER UZAYA BAĞLAN", use_container_width=True):
                secilen_index = secenekler.index(secilen_metin)
                st.session_state.aktif_cocuk_id = cocuklar[secilen_index][0]
                st.session_state.aktif_cocuk_isim = cocuklar[secilen_index][1]
                st.session_state.aktif_sayfa = "Cocuk_Panel"
                st.rerun()
            st.markdown('</div>', unsafe_allow_html=True)

def veli_giris_ekrani():
    col1, col2, col3 = st.columns([1, 1.5, 1])
    with col2:
        st.markdown('<div class="glass-box" style="text-align:center;"><h2 class="neon-text">REHBER PORTALI</h2><p style="color:#64748b;">Sisteme giriş yapın veya yeni bir rehber hesabı oluşturun.</p></div>', unsafe_allow_html=True)
        tab1, tab2 = st.tabs(["🔑 Giriş Yap", "📝 Yeni Kayıt Ol"])
        with tab1:
            g_kadi = st.text_input("Kullanıcı Adı (Giriş):", key="g_kadi")
            g_sifre = st.text_input("Şifre (Giriş):", type="password", key="g_sifre")
            st.markdown('<br>', unsafe_allow_html=True)
            if st.button("🚀 Giriş Yap", key="btn_giris"):
                if veli_giris_yap(g_kadi, g_sifre): st.session_state.veli_kadi = g_kadi; st.session_state.aktif_sayfa = "Veli_Panel"; st.rerun()
                else: st.error("Hatalı kullanıcı adı veya şifre!")
        with tab2:
            k_kadi = st.text_input("Yeni Kullanıcı Adı:", key="k_kadi")
            k_sifre = st.text_input("Yeni Şifre:", type="password", key="k_sifre")
            st.markdown('<br>', unsafe_allow_html=True)
            if st.button("💾 Hesabı Oluştur", key="btn_kayit"):
                if k_kadi and k_sifre:
                    if veli_kaydol(k_kadi, k_sifre): st.success("Hesabınız başarıyla oluşturuldu! 'Giriş Yap' sekmesinden sisteme girebilirsiniz.")
                    else: st.error("Bu kullanıcı adı zaten alınmış, lütfen başka bir tane seçin.")
                else: st.warning("Kullanıcı adı ve şifre boş bırakılamaz.")

# ==============================================================================
# SAYFA 3: VELİ PANELİ (TAKVİM & EKSİKSİZ BERAT KODU DAHİL)
# ==============================================================================
def veli_panel_ekrani():
    if not st.session_state.get("veli_kadi"):
        st.error("Lütfen önce giriş yapın.")
        if st.button("Geri Dön"): st.session_state.aktif_sayfa = "Ana Sayfa"; st.rerun()
        return

    st.markdown(f'<h1 class="neon-text">Rehberlik Köşesi</h1><p style="color:#64748b; font-size:1.1rem; margin-top:-15px; margin-bottom:30px;">Hoş Geldiniz, {st.session_state.veli_kadi}</p>', unsafe_allow_html=True)
    t1, t2, t3, t4, t5, t6, t7 = st.tabs(["Görev Onay", "Serbest Rapor", "Gelişim Matrisi", "Berat Tasarla", "Sistem Kayıt", "AI Pedagog", "📅 Takvim & Ajanda"])

    with t1:
        st.markdown('<div class="glass-box"><h3>Bekleyen Gözlemler</h3><p style="color:#64748b;">Çocukların gönderdiği görevleri onaylayın.</p></div>', unsafe_allow_html=True)
        bekleyenler = onay_bekleyenleri_getir(st.session_state.veli_kadi)
        if not bekleyenler: st.info("Şu an onay bekleyen bir görev bulunmamaktadır.")
        else:
            for kayit in bekleyenler:
                islem_id, cocuk_isim, gorev_adi, cevap, c_id, _ = kayit
                with st.expander(f"📌 {cocuk_isim} - {gorev_adi}", expanded=True):
                    st.markdown(f"<div style='background:#f1f5f9; padding:15px; border-radius:10px;'><b>Çocuğun İfadesi:</b><br>{cevap}</div><br>", unsafe_allow_html=True)
                    yuklenen_foto = st.file_uploader(f"Kanıt Fotoğrafı Yükle ({cocuk_isim})", type=['png', 'jpg', 'jpeg'], key=f"foto_{islem_id}")
                    veli_degerlendirmesi = st.text_area("Rehber Notunuzu Ekleyin:", key=f"not_{islem_id}")
                    col_onay, col_word = st.columns(2)
                    with col_onay:
                        if st.button("✅ Görevi Onayla", key=f"btn_onay_{islem_id}"): veri_onayla(islem_id, c_id); st.rerun()
                    with col_word:
                        if yuklenen_foto and veli_degerlendirmesi:
                            word_dosyasi = word_raporu_olustur(cocuk_isim, gorev_adi, veli_degerlendirmesi, yuklenen_foto.getvalue())
                            st.download_button("📄 Word Raporunu İndir", data=word_dosyasi, file_name=f"{cocuk_isim}_Rapor.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"dl_{islem_id}")

    with t2:
        st.markdown('<div class="glass-box"><h3>Serbest Fotoğraflı Rapor Al</h3></div>', unsafe_allow_html=True)
        cocuklar = cocuklari_getir(st.session_state.veli_kadi)
        if cocuklar:
            secilen_cocuk = st.selectbox("Çocuğu Seçin:", [c[1] for c in cocuklar])
            s_gorev_adi = st.text_input("Faaliyet Adı:")
            s_veli_notu = st.text_area("Gözlem Notu:")
            s_foto = st.file_uploader("Görsel Yükle:", type=['png', 'jpg', 'jpeg'])
            if s_foto and s_veli_notu and s_gorev_adi:
                s_word_dosyasi = word_raporu_olustur(secilen_cocuk, s_gorev_adi, s_veli_notu, s_foto.getvalue())
                st.download_button("📄 Raporu İndir", data=s_word_dosyasi, file_name=f"{secilen_cocuk}_Serbest_Rapor.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="dl_serbest")

    with t3:
        st.markdown('<div class="glass-box"><h3>Gelişim Matrisi ve Analiz Grafiği</h3></div>', unsafe_allow_html=True)
        cocuklar = cocuklari_getir(st.session_state.veli_kadi)
        if cocuklar:
            df = pd.DataFrame(cocuklar, columns=["ID", "İsim", "Mevcut Adım"])
            fig = px.bar(df, x="İsim", y="Mevcut Adım", text="Mevcut Adım", color="Mevcut Adım", color_continuous_scale=["#c7d2fe", "#6366f1", "#312e81"])
            fig.update_layout(plot_bgcolor="rgba(0,0,0,0)", paper_bgcolor="rgba(0,0,0,0)", font=dict(family="Plus Jakarta Sans, sans-serif", color="#334155", size=14), margin=dict(t=20, b=20, l=20, r=20))
            fig.update_traces(textfont_size=14, textposition="outside", cliponaxis=False)
            st.plotly_chart(fig, use_container_width=True)
            for cid, isim, adim in cocuklar:
                ilerleme = int((adim / max(MUFREDAT.keys())) * 100) if MUFREDAT else 0
                st.markdown(f"<div style='font-weight:600; color:#0f172a;'>{isim} <span style='color:#64748b; font-weight:400;'>- Aşama {adim} ({ilerleme}%)</span></div>", unsafe_allow_html=True)
                st.progress(min(ilerleme, 100))
        st.markdown("<hr>", unsafe_allow_html=True)
        st.markdown("#### 🎯 Yeni Görev Ata")
        if cocuklar:
            col_c, col_g = st.columns(2)
            with col_c: sec_isim = st.selectbox("Görev Atanacak Çocuk:", [c[1] for c in cocuklar], key="g_isim"); sec_id = next(c[0] for c in cocuklar if c[1] == sec_isim)
            with col_g: muf_sec = [f"Adım {a}: {i['faz']} ➔ {i['alt_seviye']}" for a, i in MUFREDAT.items()]; sec_muf = st.selectbox("Aşama:", muf_sec); y_adim = int(sec_muf.split(":")[0].replace("Adım ", ""))
            if st.button("🚀 Görevi Ata"): cocuk_adim_guncelle(sec_id, y_adim); st.rerun()

    with t4:
        st.markdown('<div class="glass-box"><h3>Özel Berat / Sertifika Tasarla</h3><p style="color:#64748b;">Görsel olarak harika sertifikalar oluşturun, bilgisayarınıza indirin ve yazdırın.</p></div>', unsafe_allow_html=True)
        cocuklar = cocuklari_getir(st.session_state.veli_kadi)
        if not cocuklar:
            st.warning("Önce 'Sisteme Kayıt' sekmesinden bir çocuk eklemelisiniz.")
        else:
            BERAT_TEMALARI = {
                "1. Antik Parşömen (Tarihi/Klasik)": {"bg": "#fefce8", "border": "#92400e", "text": "#451a03"},
                "2. Siber Uzay (Fütüristik/Mavi)": {"bg": "#eff6ff", "border": "#1e40af", "text": "#1e3a8a"},
                "3. Kökler ve Doğa (Huzur/Yeşil)": {"bg": "#f0fdf4", "border": "#166534", "text": "#14532d"},
                "4. Güneş Yıldızı (Enerji/Altın)": {"bg": "#fffbeb", "border": "#b45309", "text": "#78350f"},
                "5. Derin Okyanus (Bilgelik/Lacivert)": {"bg": "#f8fafc", "border": "#334155", "text": "#0f172a"},
                "6. Neon Işıklar (Canlı/Pembe)": {"bg": "#fdf2f8", "border": "#be185d", "text": "#831843"},
                "7. Zarif Minimal (Sade/Siyah-Beyaz)": {"bg": "#ffffff", "border": "#475569", "text": "#1e293b"}
            }

            col_sol, col_sag = st.columns([1, 1])
            with col_sol:
                secilen_isim = st.selectbox("Beratı Kazanacak Çocuğu Seçin:", [c[1] for c in cocuklar], key="berat_cocuk")
                secilen_id = next(c[0] for c in cocuklar if c[1] == secilen_isim)
                veli_isim = st.text_input("Anne/Baba Adı Soyadı:", placeholder="Örn: Ahmet Yılmaz")
                faz_adi = st.selectbox("Hangi Faz İçin Veriliyor?", [
                    "KÖKLER FAZI", "BAĞLAR FAZI", "PUSULA FAZI", 
                    "AYNALAR FAZI", "ÇARKLAR FAZI", "KÖPRÜLER FAZI", "KANATLAR FAZI", "ÖZEL BAŞARI BERATI"
                ])
                secilen_tema = st.selectbox("Sertifika Tasarımı (7 Farklı Tema):", list(BERAT_TEMALARI.keys()))
            with col_sag:
                st.info("💡 **Nasıl Çıktı Alınır?**\nBeratınızı oluşturduktan sonra 'İndir' butonuna basarak bilgisayarınıza kaydedin. Tarayıcıda açtıktan sonra klavyeden **Ctrl+P** tuşlarına basarak A4 yatay olarak yazdırabilirsiniz.")
                olustur_buton = st.button("🎨 Beratı Ekranda Çiz ve Oluştur")

            if olustur_buton:
                if not veli_isim:
                    st.error("Lütfen imza alanı için Anne/Baba adını giriniz.")
                else:
                    tema = BERAT_TEMALARI[secilen_tema]
                    html_icerik = f"""
                    <html>
                    <head>
                    <meta charset="utf-8">
                    <link href="https://fonts.googleapis.com/css2?family=Alex+Brush&family=Cinzel:wght@600;700&family=Montserrat:wght@400;500;600&display=swap" rel="stylesheet">
                    <style>
                        body {{ margin: 0; padding: 0; display: flex; justify-content: center; background: transparent; }}
                        .certificate {{ width: 1000px; height: 707px; background: {tema['bg']}; padding: 35px; box-sizing: border-box; position: relative; }}
                        .border-outer {{ width: 100%; height: 100%; border: 2px solid {tema['border']}; padding: 8px; box-sizing: border-box; }}
                        .border-inner {{ width: 100%; height: 100%; border: 4px solid {tema['border']}; padding: 50px 70px; box-sizing: border-box; position: relative; text-align: center; }}
                        .corner {{ position: absolute; width: 15px; height: 15px; background: {tema['bg']}; border: 3px solid {tema['border']}; border-radius: 50%; }}
                        .corner-tl {{ top: -10px; left: -10px; }}
                        .corner-tr {{ top: -10px; right: -10px; }}
                        .corner-bl {{ bottom: -10px; left: -10px; }}
                        .corner-br {{ bottom: -10px; right: -10px; }}
                        .header {{ font-family: 'Montserrat', sans-serif; font-size: 16px; font-weight: 500; color: {tema['text']}; text-transform: uppercase; letter-spacing: 2px; margin-bottom: 10px; }}
                        .title {{ font-family: 'Cinzel', serif; font-size: 52px; font-weight: 700; color: {tema['border']}; margin-bottom: 30px; letter-spacing: 1px; }}
                        .salutation {{ font-family: 'Montserrat', sans-serif; font-size: 18px; font-weight: 600; color: {tema['text']}; margin-bottom: 5px; }}
                        .name {{ font-family: 'Alex Brush', cursive; font-size: 85px; color: {tema['border']}; margin: 0 0 25px 0; line-height: 1.1; }}
                        .description {{ font-family: 'Montserrat', sans-serif; font-size: 16px; line-height: 1.8; color: {tema['text']}; margin: 0 auto 40px auto; max-width: 85%; }}
                        .footer {{ display: flex; justify-content: space-between; align-items: flex-end; position: absolute; bottom: 50px; left: 70px; right: 70px; }}
                        .signature-block, .qyvam-logo-block {{ text-align: center; width: 220px; position: relative; z-index: 10; }}
                        .sign-name {{ font-family: 'Montserrat', sans-serif; font-size: 16px; font-weight: 600; color: {tema['text']}; }}
                        .sign-title {{ font-family: 'Montserrat', sans-serif; font-size: 13px; color: {tema['text']}; }}
                        .sign-space {{ height: 50px; }}
                        .sign-line {{ border-top: 1px solid {tema['text']}; width: 100%; margin: 0 auto; }}
                        .seal {{ position: absolute; bottom: 30px; left: 50%; transform: translateX(-50%); opacity: 0.15; width: 140px; z-index: 1; }}
                        .qyvam-brand {{ font-family: 'Cinzel', serif; font-size: 26px; font-weight: 700; color: {tema['border']}; letter-spacing: 3px; margin-top: 25px; margin-bottom: 5px; }}
                        .qyvam-sub {{ font-family: 'Montserrat', sans-serif; font-size: 10px; letter-spacing: 1px; color: {tema['text']}; text-transform: uppercase; }}
                    </style>
                    </head>
                    <body>
                        <div class="certificate">
                            <div class="border-outer">
                                <div class="border-inner">
                                    <div class="corner corner-tl"></div>
                                    <div class="corner corner-tr"></div>
                                    <div class="corner corner-bl"></div>
                                    <div class="corner corner-br"></div>
                                    <div class="header">Qyvam Eğitim Sistemi</div>
                                    <div class="title">{faz_adi}</div>
                                    <div class="salutation">Sevgili oğlum/kızım ;</div>
                                    <div class="name">{secilen_isim}</div>
                                    <div class="description">Bu bir sertifikadan daha fazlası; birlikte geçirdiğimiz harika anların bir hatırası!<br><strong>{faz_adi}</strong> yolculuğumuza ortak olduğun ve bu güzel deneyimi beraber paylaştığımız için kalpten teşekkür ederim. Başarılarının devamını dilerim.</div>
                                    <div class="footer">
                                        <div class="signature-block">
                                            <div class="sign-name">{veli_isim}</div>
                                            <div class="sign-title">Rehber Veli</div>
                                            <div class="sign-space"></div>
                                            <div class="sign-line"></div>
                                        </div>
                                        <div class="seal">
                                            <svg viewBox="0 0 100 100" xmlns="http://www.w3.org/2000/svg">
                                                <path d="M35 60 L20 100 L38 90 L50 100 Z" fill="{tema['border']}"/>
                                                <path d="M65 60 L80 100 L62 90 L50 100 Z" fill="{tema['border']}"/>
                                                <path d="M50 5 L56 16 L68 15 L70 27 L82 31 L77 42 L85 51 L75 58 L77 70 L65 72 L59 82 L50 74 L41 82 L35 72 L23 70 L25 58 L15 51 L23 42 L18 31 L30 27 L32 15 L44 16 Z" fill="{tema['border']}"/>
                                                <circle cx="50" cy="43" r="24" fill="{tema['bg']}"/>
                                                <polygon points="50,26 56,38 69,39 59,48 62,60 50,54 38,60 41,48 31,39 44,38" fill="{tema['border']}"/>
                                            </svg>
                                        </div>
                                        <div class="qyvam-logo-block">
                                            <div class="qyvam-brand">✧ QYVAM ✧</div>
                                            <div class="qyvam-sub">Siber Uzay & Şahsiyet İnşası</div>
                                            <div class="sign-line" style="margin-top: 15px;"></div>
                                        </div>
                                    </div>
                                </div>
                            </div>
                        </div>
                    </body>
                    </html>
                    """
                    st.markdown("### 📜 Tasarım Önizlemesi")
                    import streamlit.components.v1 as components
                    components.html(html_icerik, height=750)
                    st.download_button("📥 Bu Beratı İndir (PDF / Çıktı İçin)", data=html_icerik, file_name=f"Qyvam_Berat_{secilen_isim}.html", mime="text/html")
                    ozel_berat_ekle(secilen_id, faz_adi, "Veli tarafından özel tasarım berat takdim edildi.")
                    st.success(f"Berat tasarımı hazırlandı ve {secilen_isim} isimli çocuğun profiline işlendi!")

    with t5:
        st.markdown('<div class="glass-box"><h3>Sistem Kayıt Yönetimi</h3></div>', unsafe_allow_html=True)
        col_e, col_c = st.columns(2)
        with col_e:
            yeni_isim = st.text_input("Çocuğun İsmi:")
            if st.button("Sisteme Ekle"): cocuk_ekle(yeni_isim.strip(), st.session_state.veli_kadi); st.rerun()
        with col_c:
            cocuklar = cocuklari_getir(st.session_state.veli_kadi)
            if cocuklar:
                sil_isim = st.selectbox("Silinecek Çocuk:", [c[1] for c in cocuklar])
                if st.button("❌ Profili Sil"): sil_id = next(c[0] for c in cocuklar if c[1] == sil_isim); cocuk_sil(sil_id); st.rerun()

    with t6:
        st.markdown('<div class="glass-box"><h3>Qyvam AI Pedagog</h3></div>', unsafe_allow_html=True)
        veli_sorusu = st.text_area("Pedagojik Danışmanınıza Sorun:")
        if st.button("🤖 Danış"): st.info(ai_cevap_uret(veli_sorusu, 1, rol="veli"))

    with t7:
        st.markdown('<div class="glass-box"><h3>📅 Takvim & Ajanda</h3><p style="color:#64748b;">Çocuklarınız için ileri tarihli hedefler, geziler veya hatırlatıcılar ekleyin.</p></div>', unsafe_allow_html=True)
        cocuklar = cocuklari_getir(st.session_state.veli_kadi)
        if not cocuklar:
            st.warning("Ajanda oluşturmak için önce bir çocuk eklemelisiniz.")
        else:
            col_form, col_liste = st.columns([1, 1.5])
            with col_form:
                st.markdown("#### 📝 Yeni Etkinlik Ekle")
                t_isim = st.selectbox("Kimin İçin?", [c[1] for c in cocuklar])
                t_id = next(c[0] for c in cocuklar if c[1] == t_isim)
                t_tarih = st.date_input("Tarih Seçiniz:")
                t_baslik = st.text_input("Etkinlik Başlığı (Örn: Doğa Yürüyüşü)")
                t_detay = st.text_area("Kısa Detay (İsteğe Bağlı)")
                if st.button("💾 Takvime Kaydet", key="btn_takvim_ekle"):
                    if t_baslik:
                        ajanda_ekle(st.session_state.veli_kadi, t_id, t_tarih, t_baslik, t_detay)
                        st.success("Etkinlik başarıyla takvime işlendi!")
                        st.rerun()
                    else: st.error("Lütfen bir etkinlik başlığı girin.")
            with col_liste:
                st.markdown("#### 📌 Yaklaşan Etkinlikler")
                ajanda_verileri = ajanda_getir(st.session_state.veli_kadi)
                if not ajanda_verileri: st.info("Yaklaşan herhangi bir etkinlik bulunmuyor.")
                else:
                    for kayit in ajanda_verileri:
                        k_id, k_isim, k_tarih, k_baslik, k_detay = kayit
                        t_obj = datetime.strptime(k_tarih, "%Y-%m-%d")
                        t_str = t_obj.strftime("%d.%m.%Y")
                        st.markdown(f'''<div class="agenda-card"><div class="agenda-date">📅 {t_str} - 👤 {k_isim}</div><div class="agenda-title">{k_baslik}</div><div class="agenda-detail">{k_detay}</div></div>''', unsafe_allow_html=True)
                        if st.button("🗑️ Sil", key=f"sil_t_{k_id}"): ajanda_sil(k_id); st.rerun()

# ==============================================================================
# SAYFA 4: ÇOCUK (DİJİTAL İKİZ) VERİ GİRİŞ PANELİ (BİLDİRİM EKLENDİ)
# ==============================================================================
def cocuk_panel_ekrani():
    if not st.session_state.aktif_cocuk_id:
        st.session_state.aktif_sayfa = "Ana Sayfa"
        st.rerun()
        
    bilgi = cocuk_bilgisi_getir(st.session_state.aktif_cocuk_id)
    if not bilgi: return
    isim, mevcut_adim = bilgi
    adim_bilgisi = MUFREDAT.get(mevcut_adim, {"faz": "Zirve", "ust_seviye": "Tamamlandı", "alt_seviye": "Tebrikler", "varsayilan_gorev": "Tüm adımları başarıyla tamamladın.", "varsayilan_tefekkur": "Bu yolculuk sana ne kattı?"})
    
    st.markdown(f'<h1 class="neon-text" style="font-size: 2.5rem; text-align:left;">Hoş Geldin, {isim.title()}!</h1>', unsafe_allow_html=True)
    
    # YAKLAŞAN ETKİNLİK BİLDİRİMİ
    ajanda_kayitlari = cocuk_ajanda_getir(st.session_state.aktif_cocuk_id)
    if ajanda_kayitlari:
        yakin_tarih, y_baslik, y_detay = ajanda_kayitlari[0]
        st.markdown(f'''<div class="neon-alert"><b style="font-size:1.1rem;">🔔 Yaklaşan Bir Maceran Var! ({datetime.strptime(yakin_tarih, "%Y-%m-%d").strftime("%d.%m.%Y")})</b><br>{y_baslik} - {y_detay}</div>''', unsafe_allow_html=True)
    
    st.markdown('<hr style="margin-top:-10px;">', unsafe_allow_html=True)
    
    col_img, col_hud = st.columns([1, 4])
    with col_img:
        mevcut_videolar = [dosya for dosya in os.listdir() if dosya.startswith("qyman") and dosya.endswith(".mp4")]
        if mevcut_videolar: st.video(random.choice(mevcut_videolar), autoplay=True, loop=True, muted=True)
        elif os.path.exists("qyman.png"): st.image("qyman.png", use_container_width=True)
    with col_hud:
        st.markdown(f'<div class="qyman-hud"><b>🤖 Qyman Diyor ki:</b><br>Günün görevini tamamlayıp yıldızları toplamaya ne dersin?</div>', unsafe_allow_html=True)

    t1, t2, t3 = st.tabs(["📍 Bugünün Görevi", "🏆 Kazandığım Beratlar", "💬 Qyman'a Soru Sor"])

    with t1:
        if bekleyen_gorev_kontrol(st.session_state.aktif_cocuk_id):
            st.markdown('<div class="glass-box"><h3 class="neon-text" style="color:#fca5a5;">Görev Onay Bekliyor ⏳</h3></div>', unsafe_allow_html=True)
        else:
            st.markdown(f'''<div class="glass-box"><p style="color:#64748b; font-size:1.1rem; font-weight:600; margin-bottom:5px;">Aşama {mevcut_adim} : {adim_bilgisi["alt_seviye"]}</p><h2 class="neon-text" style="margin-top:0;">{adim_bilgisi["varsayilan_gorev"]}</h2><hr style="border-color: rgba(99, 102, 241, 0.2);"><h4 style="color:#10b981;">Düşünme Vakti: {adim_bilgisi["varsayilan_tefekkur"]}</h4></div>''', unsafe_allow_html=True)
            cocuk_cevabi = st.text_area("Cevabını Buraya Yazabilirsin:", height=100)
            if st.button("✨ Görevimi Tamamladım, Gönder!"):
                if cocuk_cevabi.strip() != "": cocuk_veri_gonder(st.session_state.aktif_cocuk_id, f"[ ADIM {mevcut_adim} ]: {adim_bilgisi['alt_seviye']}", cocuk_cevabi); st.rerun()

    with t2:
        st.markdown('<div class="glass-box"><h3>Kazandığın Başarı Beratları</h3></div>', unsafe_allow_html=True)
        tum_sertifikalar = (kazanilan_sertifikalari_hesapla(mevcut_adim) or []) + (ozel_beratlari_getir(st.session_state.aktif_cocuk_id) or [])
        if not tum_sertifikalar: st.info("Henüz bir berat kazanmadın.")
        else:
            for s_adi, s_aciklama in tum_sertifikalar: st.markdown(f'<div class="glass-task"><h3 style="color:#db2777; margin-top:0;">🏆 {s_adi}</h3><p style="margin-bottom:0;">{s_aciklama}</p></div>', unsafe_allow_html=True)

    with t3:
        st.markdown('<div class="glass-box"><h3>Qyman ile Konuş</h3></div>', unsafe_allow_html=True)
        cocuk_sorusu = st.text_input("Sorunu Yaz:")
        if st.button("🤖 Qyman'a Gönder") and cocuk_sorusu:
            with st.spinner("Qyman düşünüyor..."): st.success(ai_cevap_uret(cocuk_sorusu, mevcut_adim, rol="cocuk", cocuk_isim=isim))

# ==============================================================================
# ROUTER & SOL RADAR
# ==============================================================================
ust_komuta_merkezi()
sol_radar_olustur()

if st.session_state.aktif_sayfa == "Ana Sayfa": ana_karsilama_ekrani()
elif st.session_state.aktif_sayfa == "Veli_Giris": veli_giris_ekrani()
elif st.session_state.aktif_sayfa == "Veli_Panel": veli_panel_ekrani()
elif st.session_state.aktif_sayfa == "Cocuk_Panel": cocuk_panel_ekrani()

# ==============================================================================
# ALT BİLGİ (FOOTER) MODÜLÜ (TAM EKRAN GENİŞLİĞİNDE)
# ==============================================================================
import base64

def get_base64_image(image_path):
    try:
        with open(image_path, "rb") as img_file:
            return base64.b64encode(img_file.read()).decode()
    except:
        return None

def sayfa_altbilgisi_olustur():
    logo_b64 = get_base64_image("logo.jpg") 
    if logo_b64:
        logo_html = f'<img src="data:image/jpeg;base64,{logo_b64}" class="footer-logo-img" alt="Qyvam Logo">'
    else:
        logo_html = '<h2 style="color:#4f46e5; font-family:\'Outfit\', sans-serif; margin-bottom:15px; font-weight:900; letter-spacing:1px;">✧ QYVAM ✧</h2>'

    st.markdown(f"""
    <style>
    /* FOOTER'I TAM EKRAN (FULL WIDTH) YAPAN SİHİRLİ KOD */
    .qyvam-footer {{
        width: 100vw;
        position: relative;
        left: 50%;
        right: 50%;
        margin-left: -50vw;
        margin-right: -50vw;
        background-color: #f8fafc;
        border-top: 1px solid #e2e8f0;
        padding: 60px 0 20px 0; /* Sağ ve soldaki paddingi sıfırladık, içerik container'ı hizalayacak */
        margin-top: 80px;
        font-family: 'Plus Jakarta Sans', sans-serif;
        color: #475569;
    }}
    /* İçeriklerin ortada hizalı durmasını sağlayan alan */
    .footer-container {{
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(240px, 1fr));
        gap: 50px;
        max-width: 1200px; /* İçerik çok fazla açılmasın diye limitliyoruz */
        margin: 0 auto;
        padding: 0 40px;
    }}
    .footer-col h4 {{
        color: #0f172a;
        font-family: 'Outfit', sans-serif;
        font-weight: 700;
        font-size: 1.2rem;
        margin-bottom: 25px;
    }}
    .footer-col ul {{ list-style: none; padding: 0; margin: 0; }}
    .footer-col ul li {{ margin-bottom: 12px; }}
    .footer-col a {{ text-decoration: none; color: #64748b; font-size: 0.95rem; transition: color 0.2s; }}
    .footer-col a:hover {{ color: #4f46e5; font-weight: 600; }}
    .footer-logo-img {{ max-width: 170px; margin-bottom: 15px; border-radius: 8px; }}
    .footer-slogan {{ font-size: 1rem; font-weight: 700; color: #4f46e5; margin-bottom: 15px; }}
    
    details.legal-box {{ margin-bottom: 10px; background: #ffffff; border: 1px solid #e2e8f0; border-radius: 8px; overflow: hidden; }}
    details.legal-box summary {{ padding: 10px 15px; font-size: 0.9rem; font-weight: 600; cursor: pointer; color: #334155; outline: none; transition: background 0.2s; }}
    details.legal-box summary:hover {{ background: #f1f5f9; color: #4f46e5; }}
    details.legal-box p {{ padding: 15px; font-size: 0.85rem; color: #64748b; margin: 0; border-top: 1px solid #e2e8f0; background: #f8fafc; line-height: 1.6; }}
    
    .social-links a {{ display: inline-block; margin-right: 12px; font-size: 0.9rem; font-weight: 600; background: #e2e8f0; padding: 8px 15px; border-radius: 20px; color: #475569; }}
    .social-links a:hover {{ background: #4f46e5; color: white; }}
    
    .footer-bottom {{ text-align: center; padding-top: 30px; margin-top: 60px; border-top: 1px solid #e2e8f0; font-size: 0.9rem; color: #94a3b8; font-weight: 500; width: 100%; }}
    </style>
    
    <div class="qyvam-footer">
        <div class="footer-container">
            <div class="footer-col">
                {logo_html}
                <div class="footer-slogan">Kökü Değerlerimizde, Zirvesi Şahsiyette.</div>
                <div style="font-size:0.95rem; line-height:1.8; color:#64748b;">
                    📍 Ankara, Türkiye<br>
                    ☎ +90 XXX XXX XX XX<br>
                    ✉ <a href="mailto:dedemkorkut90@gmail.com" style="color:#4f46e5; font-weight:600;">dedemkorkut90@gmail.com</a>
                </div>
            </div>
            <div class="footer-col">
                <h4>Hızlı Bağlantılar</h4>
                <ul>
                    <li><a href="#">🏠 Ana Sayfa</a></li>
                    <li><a href="#">📖 Hakkımızda</a></li>
                    <li><a href="https://mtod.mebnet.net/sites/default/files/Bili%C5%9Fsel%20Geli%C5%9Fim.pdf" target="_blank">📚 Hizmetler (Bilişsel Gelişim)</a></li>
                    <li><a href="https://mtod.mebnet.net/sites/default/files/Bili%C5%9Fsel%20Geli%C5%9Fim.pdf" target="_blank">📚 Hizmetler (Eğitim Modeli)</a></li>
                    <li><a href="#yakinda" title="Blog sayfamız çok yakında sizlerle! Bizi takip etmeye devam edin.">📰 Blog <i>(Çok Yakında)</i></a></li>
                </ul>
            </div>
            <div class="footer-col">
                <h4>Destek & SSS</h4>
                <details class="legal-box">
                    <summary>❓ Qyvam Sistemi Nedir?</summary>
                    <p>Çocukların bilişsel, ahlaki ve sosyal gelişimini müfredat tabanlı görevlerle destekleyen dijital bir şahsiyet inşası ve gelişim takip ekosistemidir.</p>
                </details>
                <details class="legal-box">
                    <summary>❓ Verilerimiz Güvende mi?</summary>
                    <p>Evet, sistemdeki verileriniz 3. şahıslara aktarılmaz. Bilgiler şifrelenmiş olarak cihazınızda KVKK standartlarına uygun şekilde korunur.</p>
                </details>
                <details class="legal-box">
                    <summary>❓ Sistemi Kimler Kullanabilir?</summary>
                    <p>Sistem, gelişim sürecini yöneten rehberler (veliler, öğretmenler, eğitmenler) ve onların gözetimindeki çocuklar için tasarlanmıştır.</p>
                </details>
                <ul style="margin-top:15px;">
                    <li><a href="mailto:dedemkorkut90@gmail.com">🎧 Yardım Merkezi (İletişim)</a></li>
                </ul>
            </div>
            <div class="footer-col">
                <h4>Yasal Bilgiler</h4>
                <details class="legal-box">
                    <summary>🛡️ Gizlilik Politikası</summary>
                    <p>Qyvam, kullanıcıların kişisel verilerini yetkisiz erişime karşı korumayı taahhüt eder. Sistemdeki değerlendirmeler ve verileriniz sadece kendi profilinizde saklanır ve ticari amaçlarla asla paylaşılamaz.</p>
                </details>
                <details class="legal-box">
                    <summary>🍪 Çerez Politikası</summary>
                    <p>Platformumuz, oturumunuzu güvenli tutmak ve cihazlar arası senkronizasyonu sağlamak için temel düzeyde teknik (zorunlu) çerezler kullanır. Reklam veya pazarlama çerezi barındırmaz.</p>
                </details>
                <details class="legal-box">
                    <summary>📜 Kullanım Şartları</summary>
                    <p>Bu platform pedagojik eğitim amaçlıdır. Kullanıcılar, çocukların psikolojik gelişimini destekleyen etik kurallara uymakla yükümlüdür. Görev onayları tamamen rehberin inisiyatifindedir.</p>
                </details>
                <details class="legal-box">
                    <summary>🔐 KVKK Aydınlatma Metni</summary>
                    <p>6698 sayılı yasa kapsamında işlenen veriler (isim, ilerleme durumu, görev metinleri), yalnızca eğitim süreçlerinin takibi ve analizi için kullanılmaktadır. Veri sahibi dilediği an profilini ve tüm verilerini kalıcı olarak silebilir.</p>
                </details>
                
                <h4 style="margin-top: 25px; margin-bottom:15px;">Sosyal Medya & Bülten</h4>
                <div class="social-links">
                    <a href="#ig" title="Çok yakında sizlerle! Bizi takip etmeye devam edin.">Instagram</a>
                    <a href="#x" title="Çok yakında sizlerle!">X</a>
                    <a href="#in" title="Çok yakında sizlerle!">LinkedIn</a>
                    <a href="#yt" title="Çok yakında sizlerle!">YouTube</a>
                </div>
                <div style="margin-top:15px; font-size:0.9rem; color:#64748b; font-weight:600;">E-Bülten: [Çok Yakında]</div>
            </div>
        </div>
        <div class="footer-bottom">
            © 2026 QYVAM. Tüm hakları saklıdır.
        </div>
    </div>
    """, unsafe_allow_html=True)

sayfa_altbilgisi_olustur()
